"""Extract plain text from PDF, DOCX, Markdown, and TXT resumes."""

from __future__ import annotations

import re
from pathlib import Path

SUPPORTED_RESUME_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".txt", ".text"}


def extract_resume_text(path: Path) -> str:
    """Return normalized plain text for a supported resume file."""
    path = path.resolve()
    if not path.exists():
        raise FileNotFoundError(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf(path)
    elif suffix == ".docx":
        text = _read_docx(path)
    elif suffix in {".md", ".markdown", ".txt", ".text"}:
        text = path.read_text(encoding="utf-8")
    else:
        raise ValueError(
            f"Unsupported resume format '{suffix}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_RESUME_EXTENSIONS))}"
        )
    return normalize_resume_text(text)


def normalize_resume_text(text: str) -> str:
    """Repair common PDF line-wrap artifacts and normalize whitespace."""
    text = text.replace("\u2022", "•").replace("\uf0b7", "•").replace("\xa0", " ")
    # Join hyphenated wraps: "containeriza-\ntion"
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    # Join mid-word wraps only when the next line starts lowercase:
    # "improvin\ng data" -> "improving data"
    text = re.sub(r"([A-Za-z])\n([a-z])", r"\1\2", text)
    lines = []
    for line in text.splitlines():
        cleaned = re.sub(r"[ \t]+", " ", line).strip()
        lines.append(cleaned)
    out: list[str] = []
    blank = 0
    for line in lines:
        if not line:
            blank += 1
            if blank <= 1:
                out.append("")
            continue
        blank = 0
        out.append(line)
    return "\n".join(out).strip() + "\n"


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ImportError("Install pypdf to import PDF resumes: pip install pypdf") from exc
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    text = "\n".join(parts)
    if not text.strip():
        raise ValueError(f"No extractable text found in PDF: {path.name}")
    return text


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ImportError(
            "Install python-docx to import DOCX resumes: pip install python-docx"
        ) from exc
    document = Document(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    if not paragraphs:
        raise ValueError(f"No extractable text found in DOCX: {path.name}")
    return "\n".join(paragraphs) + "\n"
